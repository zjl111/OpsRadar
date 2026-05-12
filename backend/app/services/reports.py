import re
import subprocess
import tempfile
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from fastapi import HTTPException
from jinja2 import Environment, FileSystemLoader, select_autoescape
from sqlalchemy.orm import Session

from backend.app.core.config import settings
from backend.app.models import InspectionReport, Issue, IssueInsight, Task, TaskResult
from backend.app.services.analysis import environment_overview
from backend.app.services.serializers import model_to_dict


env = Environment(
    loader=FileSystemLoader(settings.project_root / "backend" / "app" / "templates"),
    autoescape=select_autoescape(["html", "xml"]),
)


def _compact_text(value: str | None, limit: int = 420) -> str:
    text = " ".join(str(value or "").strip().split())
    if len(text) <= limit:
        return text
    return f"{text[:limit].rstrip()}..."


def _first_lines(text: str | None, limit: int = 6) -> str:
    lines = [line.strip() for line in str(text or "").splitlines() if line.strip()]
    return "；".join(lines[:limit])


def _extract_current_state(output: str | None, error_message: str | None, item_snapshot: dict | None = None) -> str:
    item_name = str((item_snapshot or {}).get("name") or "")
    output_text = _compact_text(output, 520)
    error_text = _compact_text(error_message, 260)
    threshold = re.search(r"Threshold failed:\s*([-+]?\d+(?:\.\d+)?)\s*([<>]=?)\s*([-+]?\d+(?:\.\d+)?)", error_text)
    evidence = _first_lines(output, 4)
    disk_lines = [line.strip() for line in str(output or "").splitlines() if "%" in line and ("/" in line or "Filesystem" in line or "文件系统" in line)]
    percent_values = [float(value) for value in re.findall(r"(\d+(?:\.\d+)?)%", str(output or ""))]
    if disk_lines and ("磁盘" in item_name or "inode" in item_name.lower()):
        evidence = "；".join(disk_lines[:5])
    if threshold:
        current_value = f"{max(percent_values):g}%" if percent_values else threshold.group(1)
        metric = "当前最高使用率" if percent_values else "当前值"
        detail = f"{metric} {current_value}，阈值要求 {threshold.group(2)} {threshold.group(3)}。"
        return f"{detail}现场输出：{evidence}" if evidence else detail
    if output_text:
        return output_text
    return error_text or "无输出。"


def _judgement_description(result: dict) -> str:
    expected = str((result.get("item_snapshot") or {}).get("expected") or "").strip()
    error = _compact_text(result.get("error_message"), 260)
    if expected:
        return f"判定规则：{expected}。{error}" if error else f"判定规则：{expected}。"
    return error or "未配置判定规则。"


def _fallback_steps(issue: dict) -> list[str]:
    item_name = str(issue.get("item_snapshot", {}).get("name") or "")
    resource_name = str(issue.get("resource_snapshot", {}).get("name") or "目标资源")
    lower_name = item_name.lower()
    if "cpu" in lower_name or "CPU" in item_name:
        return [
            f"登录 {resource_name}，执行 top、pidstat 或 ps -eo pid,comm,%cpu --sort=-%cpu 定位高 CPU 进程。",
            "确认是否为发布、批处理、备份、压测等预期负载；非预期进程需结合日志继续排查。",
            "若持续超过阈值，考虑限流、扩容、重启异常进程或调整任务执行窗口。",
            "处理后重新执行 CPU 使用率巡检，确认指标回落到阈值以内。",
        ]
    if "内存" in item_name or "memory" in lower_name:
        return [
            f"登录 {resource_name}，执行 free -m、top 或 ps -eo pid,comm,%mem --sort=-%mem 定位内存占用来源。",
            "区分 page cache、正常业务缓存和进程泄漏；检查 OOM、应用日志和容器内存限制。",
            "对异常进程执行配置优化、重启或扩容，必要时调整容器/系统内存限制。",
            "处理后重新巡检内存占用，确认使用率低于阈值。",
        ]
    if "磁盘" in item_name or "inode" in lower_name:
        return [
            f"登录 {resource_name}，执行 df -h、df -i 和 du -xh --max-depth=1 定位空间或 inode 占用目录。",
            "优先清理过期日志、临时文件、归档包和无用镜像；业务文件需确认保留策略后再处理。",
            "若清理后仍接近阈值，扩容磁盘或调整数据目录、日志轮转与保留周期。",
            "处理后重新巡检磁盘和 inode，确认所有分区低于阈值。",
        ]
    if "负载" in item_name or "load" in lower_name:
        return [
            f"登录 {resource_name}，执行 uptime、vmstat 1、iostat -x 1 和 top 判断负载来自 CPU、IO 还是阻塞。",
            "检查最近发布、定时任务、数据库慢查询、磁盘 IO 等高峰来源。",
            "按瓶颈类型处理：CPU 扩容或限流，IO 优化存储或任务窗口，阻塞进程结合日志定位。",
            "处理后观察 1/5/15 分钟 load 是否恢复到合理范围。",
        ]
    if "进程" in item_name or "process" in lower_name:
        return [
            f"登录 {resource_name}，执行 ps、pstree、systemctl 或 docker ps 查看异常进程来源。",
            "确认是否存在进程泄漏、服务频繁拉起、僵尸进程或异常子进程堆积。",
            "修复对应服务配置或进程管理策略，必要时重启异常服务并补充监控阈值。",
            "处理后重新巡检进程数量与关键进程状态。",
        ]
    if "tcp" in lower_name or "连接" in item_name:
        return [
            f"登录 {resource_name}，执行 ss -ant、ss -s 查看 TCP 状态分布。",
            "CLOSE_WAIT 偏高时排查应用连接释放；TIME_WAIT 偏高时检查短连接和连接池配置。",
            "结合应用日志、Nginx/upstream、数据库连接池确认异常来源。",
            "处理后重新巡检 TCP 状态，确认异常连接不再持续增长。",
        ]
    return [
        "查看巡检输出中的当前值、错误信息和判定规则，确认是否为真实异常。",
        "登录目标资源核对服务状态、系统日志和最近变更。",
        "按业务影响范围处理后重新执行巡检，确认异常关闭。",
    ]


def _issue_steps(issue: dict) -> list[str]:
    insight = issue.get("insight") or {}
    steps = insight.get("steps") or []
    normalized = [str(step) for step in steps if str(step).strip()][:5]
    return normalized or _fallback_steps(issue)[:5]


def _issue_recommendation(issue: dict) -> str:
    insight = issue.get("insight") or {}
    if insight.get("recommendation"):
        return str(insight["recommendation"])
    item_name = str(issue.get("item_snapshot", {}).get("name") or "该巡检项")
    return f"根据「{item_name}」的当前输出和阈值先定位异常来源，处理后重新执行巡检验证。"


def _build_report_analysis(tasks: list[dict], issues: list[dict], summary: dict) -> dict:
    total = summary.get("total", 0) or 0
    failed = (summary.get("fail", 0) or 0) + (summary.get("exception", 0) or 0)
    success = summary.get("success", 0) or 0
    task_names = "、".join(task.get("name", "") for task in tasks if task.get("name")) or "本次巡检任务"
    if not issues:
        paragraphs = [
            f"本次巡检共执行 {total} 个检查项，成功 {success} 个，未发现失败或异常项。整体看，当前应用环境基础运行状态正常。",
            "建议继续按既定调度周期执行巡检，并将本次报告作为健康基线；后续如接入监控和日志数据，可进一步补充趋势分析与容量预测。",
        ]
        return {"paragraphs": paragraphs, "key_findings": []}

    resource_counter = Counter(str(issue.get("resource_snapshot", {}).get("name") or "-") for issue in issues)
    category_counter = Counter(str(issue.get("item_snapshot", {}).get("name") or issue.get("item_snapshot", {}).get("category") or "-") for issue in issues)
    top_resources = "、".join(f"{name}（{count}项）" for name, count in resource_counter.most_common(3))
    top_categories = "、".join(f"{name}（{count}项）" for name, count in category_counter.most_common(3))
    fail_rate = round((failed / total) * 100, 1) if total else 0
    paragraphs = [
        f"本次巡检任务「{task_names}」共执行 {total} 个检查项，成功 {success} 个，失败/异常 {failed} 个，异常占比 {fail_rate}%。问题主要集中在 {top_resources}。",
        f"从巡检项类型看，重点风险集中在 {top_categories}。建议优先处理影响面较大的主机资源与容量类问题，完成处理后重新执行巡检并生成新的报告归档。",
    ]
    key_findings = [
        {
            "resource": issue.get("resource_snapshot", {}).get("name") or "-",
            "item": issue.get("item_snapshot", {}).get("name") or "-",
            "current": _extract_current_state(issue.get("output"), issue.get("error_message"), issue.get("item_snapshot")),
            "recommendation": _issue_recommendation(issue),
        }
        for issue in issues[:5]
    ]
    return {"paragraphs": paragraphs, "key_findings": key_findings}


def _task_bundle(db: Session, task_ids: list[str]) -> tuple[list[dict], list[dict], dict]:
    tasks = db.query(Task).filter(Task.id.in_(task_ids)).order_by(Task.created_at.desc()).all()
    if not tasks:
        raise HTTPException(status_code=404, detail="No tasks found for report export")

    task_dicts: list[dict] = []
    all_results: list[dict] = []
    summary = {"total": 0, "success": 0, "fail": 0, "exception": 0}

    for task in tasks:
        environment_data = None
        if task.environment:
            environment_data = {
                "id": task.environment.id,
                "name": task.environment.name,
                "env_type": task.environment.env_type,
                "application_name": task.application.name if task.application else "",
                "overview": environment_overview(db, task.environment),
            }
        results = db.query(TaskResult).filter(TaskResult.task_id == task.id).order_by(TaskResult.id).all()
        result_dicts = []
        for result in results:
            row = {
                "id": result.id,
                "task_id": task.id,
                "task_name": task.name,
                "resource_snapshot": result.resource_snapshot,
                "item_snapshot": result.item_snapshot,
                "status": result.status,
                "output": result.output,
                "error_message": result.error_message,
                "execution_time_ms": result.execution_time_ms,
            }
            result_dicts.append(row)
            all_results.append(row)
            if result.status in summary:
                summary[result.status] += 1
            summary["total"] += 1
        task_dicts.append(
            {
                "id": task.id,
                "name": task.name,
                "status": task.status,
                "started_at": task.started_at,
                "finished_at": task.finished_at,
                "application_name": task.application.name if task.application else (task.config or {}).get("application_name", ""),
                "environment_name": task.environment.name if task.environment else (task.config or {}).get("environment_name", ""),
                "environment": environment_data,
                "results": result_dicts,
            }
        )

    issues = []
    for result in all_results:
        if result["status"] not in {"fail", "exception"}:
            continue
        issue = db.query(Issue).filter(Issue.task_result_id == result["id"]).order_by(Issue.created_at.desc()).first()
        insight = db.query(IssueInsight).filter(IssueInsight.issue_id == issue.id).one_or_none() if issue else None
        issues.append(
            {
                **result,
                "issue": model_to_dict(issue) if issue else None,
                "insight": model_to_dict(insight) if insight else None,
                "severity": (issue.severity if issue else ("High" if result["status"] == "exception" else "Medium")),
                "current_state": _extract_current_state(result.get("output"), result.get("error_message"), result.get("item_snapshot")),
                "judgement": _judgement_description(result),
                "recommendation": _issue_recommendation(
                    {
                        **result,
                        "insight": model_to_dict(insight) if insight else None,
                    }
                ),
                "steps": _issue_steps(
                    {
                        **result,
                        "insight": model_to_dict(insight) if insight else None,
                    }
                ),
            }
        )
    return task_dicts, issues, summary


def render_report_html(db: Session, task_ids: list[str]) -> str:
    tasks, issues, summary = _task_bundle(db, task_ids)
    analysis = _build_report_analysis(tasks, issues, summary)
    first_env = next((task.get("environment") for task in tasks if task.get("environment")), None)
    report_title = (
        f"{first_env['application_name']} / {first_env['name']} 巡检报告"
        if first_env and len(task_ids) == 1
        else "OpsRadar Merged Inspection Report" if len(task_ids) > 1 else "OpsRadar Inspection Report"
    )
    template = env.get_template("report.html")
    return template.render(
        title=report_title,
        company_name="Operations Center",
        environment=f"{first_env['application_name']} / {first_env['name']}" if first_env else "Production",
        generated_at=datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC"),
        summary=summary,
        analysis=analysis,
        tasks=tasks,
        issues=issues,
    )


def persist_inspection_report(db: Session, task: Task) -> InspectionReport:
    existing = db.query(InspectionReport).filter(InspectionReport.task_id == task.id).order_by(InspectionReport.created_at.desc()).first()
    settings.report_dir.mkdir(parents=True, exist_ok=True)
    tasks, issues, summary_for_analysis = _task_bundle(db, [task.id])
    analysis = _build_report_analysis(tasks, issues, summary_for_analysis)
    html_path = settings.report_dir / f"{task.id}.html"
    html_path.write_text(render_report_html(db, [task.id]), encoding="utf-8")
    summary = dict(task.summary or {})
    summary.update(
        {
            "task_id": task.id,
            "task_name": task.name,
            "status": task.status,
            "finished_at": task.finished_at.isoformat() if task.finished_at else None,
            "ai_summary": analysis["paragraphs"],
        }
    )
    report = existing or InspectionReport(task_id=task.id)
    report.application_id = task.application_id
    report.environment_id = task.environment_id
    report.status = "generated" if task.status == "finished" else task.status
    report.summary = summary
    report.html_path = str(html_path)
    db.add(report)
    db.flush()
    return report


def build_docx_report(db: Session, task_ids: list[str], output_path: Path) -> Path:
    tasks, issues, summary = _task_bundle(db, task_ids)
    analysis = _build_report_analysis(tasks, issues, summary)
    doc = Document()
    first_env = next((task.get("environment") for task in tasks if task.get("environment")), None)
    doc.add_heading(f"{first_env['application_name']} / {first_env['name']} Inspection Report" if first_env else "OpsRadar Inspection Report", 0)
    doc.add_paragraph(f"Application Environment: {first_env['application_name']} / {first_env['name']}" if first_env else "Application Environment: Not specified")
    doc.add_paragraph(f"Generated at: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}")
    doc.add_heading("AI Summary & Analysis", level=1)
    for paragraph in analysis["paragraphs"]:
        doc.add_paragraph(paragraph)
    if analysis["key_findings"]:
        for finding in analysis["key_findings"]:
            doc.add_paragraph(f"{finding['resource']} / {finding['item']}: {finding['current']} 建议：{finding['recommendation']}", style=None)
    doc.add_heading("Summary", level=1)
    summary_table = doc.add_table(rows=2, cols=4)
    summary_table.style = "Table Grid"
    for index, label in enumerate(["Success", "Failed", "Exception", "Total"]):
        summary_table.cell(0, index).text = label
    for index, value in enumerate([summary["success"], summary["fail"], summary["exception"], summary["total"]]):
        summary_table.cell(1, index).text = str(value)

    for task in tasks:
        doc.add_heading(task["name"], level=1)
        doc.add_paragraph(f"Task: {task['name']} | Status: {task['status']}")
        if task.get("environment_name"):
            doc.add_paragraph(f"Application Environment: {task.get('application_name', '')} / {task.get('environment_name', '')}")
        table = doc.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        headers = ["Resource", "Type", "Address", "Check Item", "Command", "Status", "Cost"]
        for index, header in enumerate(headers):
            table.cell(0, index).text = header
        for result in task["results"]:
            row = table.add_row().cells
            row[0].text = str(result["resource_snapshot"].get("name", ""))
            row[1].text = str(result["resource_snapshot"].get("type", ""))
            row[2].text = f"{result['resource_snapshot'].get('ip')}:{result['resource_snapshot'].get('port')}"
            row[3].text = str(result["item_snapshot"].get("name", ""))
            row[4].text = str(result["item_snapshot"].get("command", ""))
            row[5].text = result["status"]
            row[6].text = f"{result['execution_time_ms']}ms"

    doc.add_heading("Issue Details", level=1)
    if issues:
        issue_table = doc.add_table(rows=1, cols=5)
        issue_table.style = "Table Grid"
        for index, header in enumerate(["Severity", "Task", "Node", "Current State", "Recommendation"]):
            issue_table.cell(0, index).text = header
        for issue in issues:
            row = issue_table.add_row().cells
            row[0].text = "High" if issue["status"] == "exception" else "Medium"
            row[1].text = issue.get("task_name") or issue["task_id"]
            row[2].text = str(issue["resource_snapshot"].get("name", ""))
            row[3].text = issue.get("current_state") or issue["output"] or issue["error_message"]
            steps = issue.get("steps") or []
            step_text = "\n".join(f"- {step}" for step in steps)
            row[4].text = "\n".join(
                part for part in [issue.get("recommendation") or "Review resource status, validate threshold, and track remediation in OpsRadar Issues.", step_text] if part
            )
    else:
        doc.add_paragraph("No issues found.")

    doc.save(output_path)
    return output_path


def build_pdf_report(html: str, output_path: Path) -> Path:
    chrome = Path(settings.chrome_path)
    if not chrome.exists():
        raise HTTPException(status_code=501, detail="Chrome headless is not configured; PDF export is unavailable")
    with tempfile.NamedTemporaryFile("w", suffix=".html", delete=False, encoding="utf-8") as handle:
        handle.write(html)
        html_path = Path(handle.name)
    try:
        subprocess.run(
            [
                str(chrome),
                "--headless",
                "--disable-gpu",
                "--no-sandbox",
                "--allow-file-access-from-files",
                f"--print-to-pdf={output_path}",
                html_path.as_uri(),
            ],
            check=True,
            timeout=40,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
        )
    except subprocess.CalledProcessError as exc:
        raise HTTPException(status_code=500, detail=f"PDF export failed: {exc.stderr.decode('utf-8', errors='ignore')[:240]}") from exc
    finally:
        html_path.unlink(missing_ok=True)
    return output_path


def export_report(db: Session, task_ids: list[str], fmt: str) -> tuple[Path, str, str]:
    safe_id = "merged" if len(task_ids) > 1 else task_ids[0]
    fmt = fmt.lower()
    settings.report_dir.mkdir(parents=True, exist_ok=True)
    if fmt == "html":
        path = settings.report_dir / f"{safe_id}.html"
        path.write_text(render_report_html(db, task_ids), encoding="utf-8")
        return path, "text/html; charset=utf-8", path.name
    if fmt in {"doc", "docx", "docs"}:
        path = settings.report_dir / f"{safe_id}.docx"
        return build_docx_report(db, task_ids, path), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", path.name
    if fmt == "pdf":
        path = settings.report_dir / f"{safe_id}.pdf"
        return build_pdf_report(render_report_html(db, task_ids), path), "application/pdf", path.name
    raise HTTPException(status_code=400, detail="Unsupported report format")
